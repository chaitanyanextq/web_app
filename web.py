import streamlit as st
import requests
from docx import Document
from PIL import Image
from github import Github
import pandas as pd
import uuid
import base64
import io
import os

# Function to check if the URL is valid
def is_valid_url(url):
    try:
        response = requests.get(url, timeout=5)
        return response.status_code == 200
    except requests.exceptions.RequestException:
        return False

# Streamlit UI Layout
title_left, title_right = st.columns([0.7, 0.3])

with title_left:
    st.title("NextQ.ai Report Generator")
with title_right:
    try:
        logo_path = "image/Screenshot_2025-02-04.png"  # Ensure this is the correct relative path
        logo = Image.open(logo_path)
        st.image(logo, width=150)
    except Exception:
        st.warning("Logo not found!")

st.write("---")

# URL section and validation
st.subheader("Input")
url = st.text_input("Enter URL")

if st.button("Check URL"):
    if is_valid_url(url):
        st.success("The URL is valid!")
    else:
        st.error("The URL is not valid.")

# GitHub credentials (Replace with environment variable or secret storage)
GITHUB_TOKEN = st.secrets["GIT"] # Use environment variable
REPO_NAME = "chaitanyanextq/web_app"
FILE_PATH = "web_data.csv"

# Initialize GitHub
g = Github(GITHUB_TOKEN)
repo = g.get_repo(REPO_NAME)

def get_csv_content():
    """Fetch existing CSV from GitHub"""
    try:
        file = repo.get_contents(FILE_PATH)
        csv_content = base64.b64decode(file.content).decode('utf-8')
        df = pd.read_csv(io.StringIO(csv_content))
        return df, file.sha
    except Exception:
        return pd.DataFrame(columns=["id", "url"]), None  # Return empty DataFrame if file not found

def append_url(url):
    """Append a new URL with a unique ID to the CSV and upload it"""
    df, file_sha = get_csv_content()
    
    # Generate unique ID
    unique_id = str(uuid.uuid4())[:8]
    
    # Append new data
    new_data = pd.DataFrame([[unique_id, url]], columns=["id", "url"])
    df = pd.concat([df, new_data], ignore_index=True)
    
    # Convert DataFrame to CSV
    csv_data = df.to_csv(index=False)

    try:
        if file_sha:
            repo.update_file(FILE_PATH, "Updating CSV with new URL", csv_data, file_sha)
        else:
            repo.create_file(FILE_PATH, "Creating CSV with first entry", csv_data)
        return unique_id
    except Exception as e:
        st.error(f"Error updating GitHub: {e}")
        return None

# If a valid URL is provided, store and generate the report
if url:
    if st.button("Add URL to Database"):
        unique_id = append_url(url)
        if unique_id:
            st.success(f"URL added successfully with ID: {unique_id}")

if url and st.button("Generate Report"):
    st.write("Generating Report...")

    # Create a Word document
    doc = Document()
    doc.add_heading('NextQ.ai Report', level=1)
    doc.add_paragraph(f'This is the report for the provided URL:\n{url}\n\n')

    # Save the document
    report_path = "NextQ_ai_Report.docx"
    doc.save(report_path)

    st.success("Report Generated Successfully!")
    st.write("Download the report from the link below:")

    # Provide download button
    with open(report_path, "rb") as file:
        st.download_button(
            label="Download Report",
            data=file,
            file_name=report_path,
            mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
